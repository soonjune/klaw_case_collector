#-*- coding:utf-8 -*-

import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime


def get_driver():

    options = webdriver.ChromeOptions()
    # prefs = {'profile.default_content_setting_values': {'cookies': 2, 'images': 2, 'javascript': 2,
    #                                                     'plugins': 2, 'popups': 2, 'geolocation': 2,
    #                                                     'notifications': 2, 'auto_select_certificate': 2,
    #                                                     'fullscreen': 2,
    #                                                     'mouselock': 2, 'mixed_script': 2, 'media_stream': 2,
    #                                                     'media_stream_mic': 2, 'media_stream_camera': 2,
    #                                                     'protocol_handlers': 2,
    #                                                     'ppapi_broker': 2, 'automatic_downloads': 2, 'midi_sysex': 2,
    #                                                     'push_messaging': 2, 'ssl_cert_decisions': 2,
    #                                                     'metro_switch_to_desktop': 2,
    #                                                     'protected_media_identifier': 2, 'app_banner': 2,
    #                                                     'site_engagement': 2,
    #                                                     'durable_storage': 2}}
    # options.add_argument('headless')
    # options.add_experimental_option('prefs', prefs)
    # options.add_argument("start-maximized")
    # options.add_argument("disable-infobars")
    # options.add_argument("--disable-extensions")
    # options.add_argument('window-size=1920x1080')
    # options.add_argument("disable-gpu")
    # options.add_argument(
    #     "user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/61.0.3163.100 Safari/537.36")
    driver = webdriver.Chrome(path, options=options)
    return driver


path = "./chromedriver.exe"
url = "https://glaw.scourt.go.kr/wsjo/intesrch/sjo022.do"

cases = pd.read_csv("cases.csv")['판례'].to_list() #cases = pd.read_csv("cases.csv", encoding='cp949')['판례'].to_list()

document = Document()

for case in cases:
    try:
        case = case.strip()
        driver = get_driver()
        driver.implicitly_wait(3)
        driver.get(url)
        time.sleep(2)
        search_box = driver.find_element_by_name("srchw")
        search_box.send_keys(case)
        search_box.send_keys(Keys.RETURN)
        time.sleep(2)
        driver.switch_to_window(driver.window_handles[1])
        driver.get_window_position(driver.window_handles[1])
        time.sleep(0.3)

        html = driver.page_source
        soup = BeautifulSoup(html, 'html.parser')
    
        name = soup.find_all("h2")[1].text.replace(u'\xa0', u' ')
        summary = soup.find_all("p", {"class": "sub_title"})[0].text
        text = soup.find_all("div", {"class": "page"})[0].find_all("p")
        driver.quit()
    
        document.add_heading(name)
        paragraph = document.add_paragraph('')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(11)
        paragraph_format.space_before = Pt(11)
        document.add_paragraph(summary)
        document.add_paragraph('')
    
        whitespace_flag = False
        for t in text:
            if "【전문】" == t.text.replace(" ", ""):
                break

            if "【참조판례】" == t.text.replace(" ", ""):
                break
    
            if t.text.isspace():
                if whitespace_flag:
                    continue
                else:
                    whitespace_flag = True
            else:
                document.add_paragraph(t.text)
                whitespace_flag=False
    
    except:
        print(f"{case} 실패")
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')

date = datetime.today().strftime('%Y%m%d')
document.save(f"{date}.docx")
